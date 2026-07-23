"""첨부 문서 파싱 — PDF/DOCX/XLSX/TXT/MD → 텍스트, 컨텍스트 예산 클리핑 (design 2026-07-23)."""

import io
from pathlib import PurePosixPath

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md"}
MAX_ATTACHMENT_BYTES = 20 * 1024 * 1024  # 업로드 상한 20MB


class ParseError(Exception):
    """파싱 실패 — message는 사용자 표시용(내부 스택 노출 금지)."""


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader  # 무거운 import 지연 — 파싱 경로에서만

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _parse_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _parse_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("cp949")  # 사내 Windows 산출 텍스트 폴백


def parse_attachment(filename: str, data: bytes) -> str:
    """확장자 디스패치 파싱 — 실패는 전부 ParseError로 정규화(라우터가 400/상태 기록)."""
    ext = PurePosixPath(filename.lower()).suffix
    if ext not in ALLOWED_EXTENSIONS:
        raise ParseError(f"unsupported file type: {ext or filename}")
    try:
        if ext == ".pdf":
            return _parse_pdf(data)
        if ext == ".docx":
            return _parse_docx(data)
        if ext == ".xlsx":
            return _parse_xlsx(data)
        return _parse_text(data)
    except ParseError:
        raise
    except Exception as exc:  # noqa: BLE001 -- 라이브러리별 예외를 경계에서 정규화
        raise ParseError(f"failed to parse {filename}") from exc


def clip_to_budget(sections: list[tuple[str, str]], budget: int) -> str:
    """[파일명] 헤더 섹션 합본 — 총량이 예산을 넘으면 각 섹션을 균등 비율로 절단."""
    if not sections:
        return ""
    total = sum(len(text) for _, text in sections)
    parts: list[str] = []
    for name, text in sections:
        if total > budget:
            share = budget * len(text) // max(total, 1)
            text = text[:share]
        parts.append(f"[{name}]\n{text}")
    return "\n\n".join(parts)
